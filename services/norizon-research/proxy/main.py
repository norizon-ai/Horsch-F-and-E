"""
Nora API Proxy - FastAPI proxy for DeepResearch API
Forwards requests to DeepResearch backend and streams SSE responses.
Production-ready with health checks, error handling, and configurable endpoints.
"""
import os
import asyncio
import json
import logging
from typing import AsyncGenerator
from contextlib import asynccontextmanager

import httpx
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from bs4 import BeautifulSoup

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuration from environment variables
DEEPRESEARCH_URL = os.getenv("DEEPRESEARCH_URL", "http://deepresearch:5000")
PROXY_PORT = int(os.getenv("PROXY_PORT", "5002"))
REQUEST_TIMEOUT = int(os.getenv("REQUEST_TIMEOUT", "300"))  # 5 minutes for research
CORS_ORIGINS = os.getenv("CORS_ORIGINS", "*").split(",")

# DeepResearch credentials (auto-creates user if doesn't exist)
DEEPRESEARCH_USERNAME = os.getenv("DEEPRESEARCH_USERNAME", "nora-proxy")
DEEPRESEARCH_PASSWORD = os.getenv("DEEPRESEARCH_PASSWORD", "nora-proxy-2025")

# Global HTTP client
http_client: httpx.AsyncClient = None
session_cookies: dict = {}


# Request/Response Models
class QueryRequest(BaseModel):
    query: str = Field(..., min_length=1, max_length=1000)
    session_id: str = Field(..., min_length=1)
    include_web: bool = True
    max_sources: int = Field(default=5, ge=1, le=20)


class HealthResponse(BaseModel):
    status: str
    deepresearch_connected: bool
    deepresearch_url: str


async def login_to_deepresearch() -> bool:
    """Login to DeepResearch and store session cookies"""
    global session_cookies

    try:
        logger.info(f"Logging in to DeepResearch as '{DEEPRESEARCH_USERNAME}'...")

        # First get CSRF token for login
        csrf_response = await http_client.get(f"{DEEPRESEARCH_URL}/auth/csrf-token")
        if csrf_response.status_code != 200:
            logger.error(f"Failed to get CSRF token: {csrf_response.status_code}")
            return False

        csrf_token = csrf_response.json().get("csrf_token", "")

        # Perform login with form data (not JSON)
        login_response = await http_client.post(
            f"{DEEPRESEARCH_URL}/auth/login",
            data={
                "username": DEEPRESEARCH_USERNAME,
                "password": DEEPRESEARCH_PASSWORD,
                "csrf_token": csrf_token,
                "remember": "true"
            },
            headers={"Content-Type": "application/x-www-form-urlencoded"}
        )

        if login_response.status_code == 200:
            # Store session cookies
            session_cookies = dict(login_response.cookies)
            logger.info("✓ Successfully logged in to DeepResearch")
            return True
        else:
            logger.error(f"Login failed: {login_response.status_code} - {login_response.text[:200]}")
            return False

    except Exception as e:
        logger.error(f"Login error: {e}")
        return False


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize and cleanup HTTP client"""
    global http_client

    logger.info(f"Starting Nora API Proxy")
    logger.info(f"DeepResearch URL: {DEEPRESEARCH_URL}")
    logger.info(f"Proxy Port: {PROXY_PORT}")

    # Initialize HTTP client with connection pooling
    http_client = httpx.AsyncClient(
        timeout=httpx.Timeout(REQUEST_TIMEOUT, connect=10.0),
        limits=httpx.Limits(max_keepalive_connections=20, max_connections=100),
        follow_redirects=True
    )

    # Test connection to DeepResearch
    try:
        logger.info("Testing connection to DeepResearch...")
        response = await http_client.get(f"{DEEPRESEARCH_URL}/")
        # Accept any response (including redirects) as a sign of connectivity
        if response.status_code in [200, 302, 401, 404]:
            logger.info("✓ Successfully connected to DeepResearch")
        else:
            logger.warning(f"DeepResearch check returned status {response.status_code}")
    except Exception as e:
        logger.error(f"Failed to connect to DeepResearch: {e}")
        logger.warning("Proxy will start anyway, but requests will fail until DeepResearch is available")

    # Login to DeepResearch
    await login_to_deepresearch()

    yield

    # Cleanup
    logger.info("Shutting down HTTP client...")
    await http_client.aclose()


# Create FastAPI app
app = FastAPI(
    title="Nora API Proxy",
    description="Production-ready proxy for DeepResearch API with SSE streaming",
    version="2.0.0",
    lifespan=lifespan
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """
    Health check endpoint
    Returns proxy status and DeepResearch connectivity
    """
    deepresearch_connected = False

    try:
        response = await http_client.get(
            f"{DEEPRESEARCH_URL}/",
            timeout=5.0
        )
        # Accept any response (including redirects) as a sign of connectivity
        deepresearch_connected = response.status_code in [200, 302, 401, 404]
    except Exception as e:
        logger.debug(f"DeepResearch connectivity check failed: {e}")

    return HealthResponse(
        status="healthy" if deepresearch_connected else "degraded",
        deepresearch_connected=deepresearch_connected,
        deepresearch_url=DEEPRESEARCH_URL
    )


async def get_csrf_token() -> str:
    """Get CSRF token from DeepResearch"""
    try:
        response = await http_client.get(f"{DEEPRESEARCH_URL}/auth/csrf-token")
        if response.status_code == 200:
            data = response.json()
            return data.get("csrf_token", "")
        logger.warning(f"Failed to get CSRF token: {response.status_code}")
        return ""
    except Exception as e:
        logger.error(f"Error getting CSRF token: {e}")
        return ""


@app.post("/query")
async def query_stream(request: QueryRequest):
    """
    Stream research results from DeepResearch API
    1. Start research and get research_id
    2. Poll for results
    3. Stream results as SSE to client
    """
    try:
        logger.info(f"Proxying query: '{request.query[:50]}...' (session: {request.session_id})")

        # Check if we have session cookies, if not try to login
        if not session_cookies:
            logger.warning("No session cookies, attempting to login...")
            if not await login_to_deepresearch():
                raise HTTPException(status_code=503, detail="Failed to authenticate with DeepResearch")

        # Get CSRF token
        csrf_token = await get_csrf_token()
        if not csrf_token:
            raise HTTPException(status_code=503, detail="Failed to obtain CSRF token from DeepResearch")

        # Prepare headers
        headers = {
            "Content-Type": "application/json",
            "X-CSRF-Token": csrf_token
        }

        # Stream the SSE response with polling
        async def event_generator() -> AsyncGenerator[bytes, None]:
            try:
                # Step 1: Start research
                logger.info(f"Starting research for: {request.query[:50]}...")
                start_response = await http_client.post(
                    f"{DEEPRESEARCH_URL}/api/start_research",
                    json={"query": request.query},
                    headers=headers,
                    cookies=session_cookies,
                    timeout=30.0
                )

                if start_response.status_code != 200:
                    error_msg = start_response.text[:200]
                    logger.error(f"Failed to start research: {start_response.status_code} - {error_msg}")
                    error_data = f'data: {json.dumps({"type": "error", "data": {"message": "Failed to start research"}})}\n\n'
                    yield error_data.encode()
                    return

                result = start_response.json()
                research_id = result.get("research_id")

                if not research_id:
                    logger.error(f"No research_id in response: {result}")
                    error_data = f'data: {json.dumps({"type": "error", "data": {"message": "Invalid response from DeepResearch"}})}\n\n'
                    yield error_data.encode()
                    return

                logger.info(f"Research started with ID: {research_id}")

                # Send initial status
                status_data = f'data: {json.dumps({"type": "status", "data": {"message": "Research started..."}})}\n\n'
                yield status_data.encode()

                # Step 2: Poll for results
                poll_interval = 1.0  # Start with 1 second
                max_polls = 300  # 5 minutes max
                poll_count = 0

                while poll_count < max_polls:
                    await asyncio.sleep(poll_interval)
                    poll_count += 1

                    # Get research status/results
                    try:
                        results_response = await http_client.get(
                            f"{DEEPRESEARCH_URL}/api/research/{research_id}",
                            headers={"X-CSRF-Token": csrf_token},
                            cookies=session_cookies,
                            timeout=10.0
                        )

                        if results_response.status_code == 200:
                            research_data = results_response.json()
                            status = research_data.get("status", "unknown")

                            logger.debug(f"Poll {poll_count}: status={status}")

                            # Check if complete
                            if status == "completed":
                                logger.info(f"Research completed: {research_id}")

                                answer = None

                                # Try multiple methods to get the report
                                # Method 1: Try /api/report/{research_id} or /api/research/{research_id}/report
                                try:
                                    logger.info(f"Trying /api/report/{research_id}")
                                    report_response = await http_client.get(
                                        f"{DEEPRESEARCH_URL}/api/report/{research_id}",
                                        headers={"X-CSRF-Token": csrf_token},
                                        cookies=session_cookies,
                                        timeout=10.0
                                    )
                                    if report_response.status_code != 200:
                                        logger.info(f"Trying /api/research/{research_id}/report")
                                        report_response = await http_client.get(
                                            f"{DEEPRESEARCH_URL}/api/research/{research_id}/report",
                                            headers={"X-CSRF-Token": csrf_token},
                                            cookies=session_cookies,
                                            timeout=10.0
                                        )
                                    if report_response.status_code != 200:
                                        logger.info(f"Trying /results/{research_id}")
                                        report_response = await http_client.get(
                                            f"{DEEPRESEARCH_URL}/results/{research_id}",
                                            headers={"X-CSRF-Token": csrf_token},
                                            cookies=session_cookies,
                                            timeout=10.0
                                        )
                                    if report_response.status_code == 200:
                                        # Check if it's HTML (web page) or JSON
                                        content_type = report_response.headers.get("content-type", "")
                                        if "application/json" in content_type:
                                            result_json = report_response.json()
                                            answer = result_json.get("report") or result_json.get("answer") or result_json.get("content")
                                        else:
                                            # It's likely HTML, extract text from the report section
                                            html_content = report_response.text
                                            soup = BeautifulSoup(html_content, 'html.parser')

                                            # Remove navigation, headers, footers, scripts, styles
                                            for element in soup(['nav', 'header', 'footer', 'script', 'style', 'button', 'form']):
                                                element.decompose()

                                            # Try to find the main report content - be more specific
                                            report_section = (
                                                soup.find('pre', class_=lambda x: x and 'markdown' in x.lower()) or
                                                soup.find('div', class_=lambda x: x and any(term in str(x).lower() for term in ['report', 'content', 'research', 'result'])) or
                                                soup.find(id=lambda x: x and any(term in str(x).lower() for term in ['report', 'content', 'research', 'result'])) or
                                                soup.find('pre') or
                                                soup.find('article') or
                                                soup.find('main')
                                            )

                                            if report_section:
                                                # Get text with some formatting preserved
                                                answer = report_section.get_text(separator='\n', strip=True)
                                                # If it's too short (just UI elements), try finding any pre/code blocks
                                                if len(answer) < 100:
                                                    pre_blocks = soup.find_all('pre')
                                                    if pre_blocks:
                                                        answer = '\n\n'.join(block.get_text(strip=True) for block in pre_blocks)
                                            else:
                                                # Last resort: get all pre/code blocks
                                                pre_blocks = soup.find_all('pre')
                                                if pre_blocks:
                                                    answer = '\n\n'.join(block.get_text(strip=True) for block in pre_blocks)
                                                else:
                                                    # Fallback to body text (cleaned)
                                                    answer = soup.body.get_text(separator='\n', strip=True) if soup.body else soup.get_text(separator='\n', strip=True)

                                        logger.info(f"Got report from /results/, length: {len(answer) if answer else 0}")
                                        logger.info(f"Report preview (first 200 chars): {answer[:200] if answer else 'None'}")
                                except Exception as e:
                                    logger.warning(f"Failed to fetch from /results/: {e}")

                                # Method 2: Try report_path if available
                                if not answer:
                                    report_path = research_data.get("report_path")
                                    if report_path:
                                        logger.info(f"Fetching report from: {report_path}")
                                        try:
                                            report_response = await http_client.get(
                                                f"{DEEPRESEARCH_URL}{report_path}",
                                                headers={"X-CSRF-Token": csrf_token},
                                                cookies=session_cookies,
                                                timeout=10.0
                                            )
                                            if report_response.status_code == 200:
                                                answer = report_response.text
                                                logger.info(f"Got report from report_path, length: {len(answer)}")
                                        except Exception as e:
                                            logger.warning(f"Failed to fetch from report_path: {e}")

                                # Method 3: Direct fields
                                if not answer:
                                    answer = research_data.get("answer") or research_data.get("result") or research_data.get("text") or research_data.get("report")

                                if answer:
                                    # Parse sources from markdown before sending content
                                    import re
                                    sources = []

                                    # Look for markdown links in citation format: [[1]](URL) Title
                                    # or numbered format: [1] Title\n   URL: https://...
                                    link_pattern = r'\[\[(\d+)\]\]\((https?://[^\)]+)\)\s*([^\n]+)'
                                    numbered_pattern = r'\[(\d+)\]\s+([^\n]+?)\s+URL:\s+(https?://[^\s]+)'

                                    for match in re.finditer(link_pattern, answer):
                                        idx, url, title = match.groups()
                                        sources.append({
                                            "id": f"src-{idx}",
                                            "title": title.strip(),
                                            "url": url.strip(),
                                            "snippet": "",
                                            "relevance": 0.9
                                        })

                                    for match in re.finditer(numbered_pattern, answer):
                                        idx, title, url = match.groups()
                                        sources.append({
                                            "id": f"src-{idx}",
                                            "title": title.strip(),
                                            "url": url.strip(),
                                            "snippet": "",
                                            "relevance": 0.9
                                        })

                                    # Remove duplicates based on URL
                                    unique_sources = []
                                    seen_urls = set()
                                    for source in sources:
                                        if source["url"] not in seen_urls:
                                            seen_urls.add(source["url"])
                                            unique_sources.append(source)

                                    logger.info(f"Extracted {len(unique_sources)} sources from report")

                                    # Send sources first
                                    for source in unique_sources:
                                        sse_data = {"type": "source", "data": source}
                                        source_data = f'data: {json.dumps(sse_data)}\n\n'
                                        yield source_data.encode()

                                    # Then send content
                                    sse_data = {"type": "content", "data": answer}
                                    content_data = f'data: {json.dumps(sse_data)}\n\n'
                                    yield content_data.encode()
                                else:
                                    logger.error(f"No answer found after trying all methods. Research data: {research_data}")

                                # Send done
                                done_data = f'data: {json.dumps({"type": "done", "data": None})}\n\n'
                                yield done_data.encode()
                                break

                            elif status == "failed" or status == "error":
                                error_msg = research_data.get("error", "Research failed")
                                logger.error(f"Research failed: {error_msg}")
                                error_data = f'data: {json.dumps({"type": "error", "data": {"message": error_msg}})}\n\n'
                                yield error_data.encode()
                                break

                        elif results_response.status_code == 404:
                            # Research not found yet, keep polling
                            continue

                    except httpx.TimeoutException:
                        logger.warning(f"Timeout polling research {research_id}, retrying...")
                        continue
                    except Exception as poll_error:
                        logger.error(f"Error polling: {poll_error}")
                        continue

                    # Adaptive polling: slow down after initial rapid polls
                    if poll_count > 10:
                        poll_interval = 2.0

                if poll_count >= max_polls:
                    logger.warning(f"Research {research_id} timed out after {max_polls} polls")
                    timeout_data = f'data: {json.dumps({"type": "error", "data": {"message": "Research timed out"}})}\n\n'
                    yield timeout_data.encode()

                logger.info(f"Query completed (session: {request.session_id})")

            except Exception as e:
                logger.error(f"Error in event generator: {e}", exc_info=True)
                error_data = f'data: {json.dumps({"type": "error", "data": {"message": f"Internal error: {str(e)}"}})}\n\n'
                yield error_data.encode()

        return StreamingResponse(
            event_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no"
            }
        )

    except httpx.ConnectError as e:
        logger.error(f"Cannot connect to DeepResearch: {e}")
        raise HTTPException(
            status_code=503,
            detail=f"Cannot connect to DeepResearch at {DEEPRESEARCH_URL}"
        )
    except httpx.TimeoutException as e:
        logger.error(f"Request timeout: {e}")
        raise HTTPException(
            status_code=504,
            detail="DeepResearch request timed out"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Internal proxy error: {str(e)}"
        )


@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "service": "Nora API Proxy",
        "version": "2.0.0",
        "status": "running",
        "deepresearch_url": DEEPRESEARCH_URL
    }


# Export Models
class ExportSessionRequest(BaseModel):
    """Session data for export"""
    session_id: str
    title: str
    messages: list
    created_at: int
    updated_at: int


@app.post("/sessions/{session_id}/export/pdf")
async def export_session_pdf(session_id: str, session_data: ExportSessionRequest):
    """Export research session to PDF"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT
        import io
        from datetime import datetime
        from pathlib import Path

        logger.info(f"Exporting session {session_id} to PDF")

        # Create PDF in memory
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        story = []
        styles = getSampleStyleSheet()

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#FF6B35',  # Norizon orange
            spaceAfter=30,
            alignment=TA_LEFT
        )
        story.append(Paragraph(session_data.title, title_style))
        story.append(Spacer(1, 0.3*inch))

        # Session Info
        created_date = datetime.fromtimestamp(session_data.created_at / 1000).strftime('%B %d, %Y at %I:%M %p')
        info_text = f"<b>Created:</b> {created_date}<br/>"
        info_text += f"<b>Messages:</b> {len(session_data.messages)}"
        story.append(Paragraph(info_text, styles['Normal']))
        story.append(Spacer(1, 0.4*inch))

        # Conversation
        story.append(Paragraph("<b>Research Conversation</b>", styles['Heading2']))
        story.append(Spacer(1, 0.2*inch))

        # Message style
        user_style = ParagraphStyle(
            'UserMessage',
            parent=styles['Normal'],
            fontSize=11,
            textColor='#0066CC',  # Norizon blue
            leftIndent=20,
            spaceAfter=10
        )

        assistant_style = ParagraphStyle(
            'AssistantMessage',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=10,
            leading=14
        )

        for idx, message in enumerate(session_data.messages):
            role = message.get('role', 'user')
            content = message.get('content', '')

            # Role header
            role_text = "You" if role == 'user' else "Nora"
            role_color = '#0066CC' if role == 'user' else '#FF6B35'

            role_style = ParagraphStyle(
                f'Role{idx}',
                parent=styles['Normal'],
                fontSize=10,
                textColor=role_color,
                fontName='Helvetica-Bold',
                spaceAfter=4
            )

            story.append(Paragraph(role_text, role_style))

            # Message content
            # Escape HTML entities and convert markdown-style formatting
            content = content.replace('<', '&lt;').replace('>', '&gt;')
            content = content.replace('\n', '<br/>')

            msg_style = user_style if role == 'user' else assistant_style
            story.append(Paragraph(content, msg_style))
            story.append(Spacer(1, 0.2*inch))

        # Build PDF
        doc.build(story)
        buffer.seek(0)

        # Save to temp file
        filename = f"{session_data.title.replace(' ', '_')}.pdf"
        temp_path = Path(f"/tmp/research_{session_id}.pdf")

        with open(temp_path, 'wb') as f:
            f.write(buffer.getvalue())

        from fastapi.responses import FileResponse
        return FileResponse(
            path=temp_path,
            media_type="application/pdf",
            filename=filename,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

    except Exception as e:
        logger.error(f"Error exporting PDF: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to export PDF: {str(e)}")


@app.post("/sessions/{session_id}/export/docx")
async def export_session_docx(session_id: str, session_data: ExportSessionRequest):
    """Export research session to DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        from datetime import datetime
        from pathlib import Path

        logger.info(f"Exporting session {session_id} to DOCX")

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(session_data.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.color.rgb = RGBColor(255, 107, 53)  # Norizon orange

        # Session Info
        created_date = datetime.fromtimestamp(session_data.created_at / 1000).strftime('%B %d, %Y at %I:%M %p')

        info_para = doc.add_paragraph()
        info_para.add_run('Created: ').bold = True
        info_para.add_run(f"{created_date}\n")
        info_para.add_run('Messages: ').bold = True
        info_para.add_run(str(len(session_data.messages)))

        doc.add_paragraph()  # Spacer

        # Conversation
        conv_heading = doc.add_heading('Research Conversation', 1)
        for run in conv_heading.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)  # Norizon blue

        # Messages
        for message in session_data.messages:
            role = message.get('role', 'user')
            content = message.get('content', '')

            # Role header
            role_text = "You" if role == 'user' else "Nora"
            role_color = RGBColor(0, 102, 204) if role == 'user' else RGBColor(255, 107, 53)

            role_para = doc.add_paragraph()
            role_run = role_para.add_run(role_text)
            role_run.bold = True
            role_run.font.color.rgb = role_color
            role_run.font.size = Pt(10)

            # Message content
            msg_para = doc.add_paragraph(content)
            msg_para.paragraph_format.left_indent = Inches(0.3)
            msg_para.paragraph_format.space_after = Pt(12)

            for run in msg_para.runs:
                run.font.size = Pt(11)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Save to temp file
        filename = f"{session_data.title.replace(' ', '_')}.docx"
        temp_path = Path(f"/tmp/research_{session_id}.docx")

        with open(temp_path, 'wb') as f:
            f.write(buffer.getvalue())

        from fastapi.responses import FileResponse
        return FileResponse(
            path=temp_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

    except Exception as e:
        logger.error(f"Error exporting DOCX: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to export DOCX: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=PROXY_PORT,
        reload=True,
        log_level="info"
    )
